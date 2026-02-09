import os
import time
import asyncio
from pathlib import Path
from typing import Dict, Any, List, Optional
from celery import Task
from sqlmodel import Session, select
from tasks.celery_app import celery_app
from core.logger import logger
from services.embedding import generate_vector
from services.graph_extractor import get_graph_extractor
from database import engine
from models import Document, Entity, GraphEdge

class BaseTaskWithDB(Task):
    """带有数据库连接的基础任务类"""
    
    def __init__(self):
        super().__init__()
        self.db_session = None
    
    def before_start(self, task_id, args, kwargs):
        """任务开始前初始化数据库连接"""
        logger.info(f"任务 {task_id} 开始执行")
        self.db_session = Session(engine)
    
    def on_failure(self, exc, task_id, args, kwargs, einfo):
        """任务失败时的处理"""
        logger.error(f"任务 {task_id} 失败: {exc}", exc_info=True)
        if self.db_session:
            self.db_session.rollback()
    
    def after_return(self, status, retval, task_id, args, kwargs, einfo):
        """任务返回后的清理"""
        if self.db_session:
            self.db_session.close()
        logger.info(f"任务 {task_id} 完成，状态: {status}")

# ==================== 文件处理辅助函数 ====================

def _process_pdf(file_path: str) -> tuple[str, int]:
    """处理 PDF 文件"""
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            content = ""
            for page_num in range(total_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:  # 确保有内容
                    content += text + "\n\n"
                
                # 每处理10页记录一次进度
                if page_num % 10 == 0 and page_num > 0:
                    logger.debug(f"📄 已处理 {page_num+1}/{total_pages} 页")
            
            logger.info(f"✅ PDF 解析完成: {total_pages} 页")
            return content, total_pages
            
    except ImportError:
        raise ImportError("请安装 PyPDF2: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"PDF 处理失败: {e}")

def _process_docx(file_path: str) -> tuple[str, int]:
    """处理 DOCX 文件"""
    try:
        import docx
        
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        content = "\n".join(paragraphs)
        logger.info(f"✅ DOCX 解析完成: {len(paragraphs)} 段落")
        
        return content, len(paragraphs)
        
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"DOCX 处理失败: {e}")

def _process_text(file_path: str) -> tuple[str, int]:
    """处理文本文件"""
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                    content = file.read()
                
                # 检查内容是否有效
                if content and len(content) > 10:
                    lines = content.count('\n') + 1
                    logger.info(f"✅ 文本文件解析完成 ({encoding}编码): {lines} 行")
                    return content, lines
                    
            except UnicodeDecodeError:
                continue
            except Exception:
                continue
        
        # 如果所有编码都失败，使用二进制读取
        with open(file_path, 'rb') as file:
            content = file.read().decode('utf-8', errors='ignore')
        
        lines = content.count('\n') + 1
        logger.warning(f"⚠️ 使用二进制模式解析文本文件: {lines} 行")
        
        return content, lines
        
    except Exception as e:
        raise Exception(f"文本文件处理失败: {e}")

def _split_into_chunks(text: str, max_chunk_size: int = 2000, overlap: int = 200) -> list[str]:
    """将文本分割为重叠的块"""
    if not text:
        return []
    
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        end = min(start + max_chunk_size, text_length)
        
        # 如果不在段落边界，向前找合适的边界
        if end < text_length:
            # 尝试在段落边界分割
            paragraph_end = text.find('\n\n', start + int(max_chunk_size * 0.8))
            if paragraph_end != -1 and paragraph_end < end + 500:
                end = paragraph_end + 2  # 包括换行符
            else:
                # 尝试在句子边界分割
                sentence_end = text.find('. ', start + int(max_chunk_size * 0.8))
                if sentence_end != -1 and sentence_end < end + 300:
                    end = sentence_end + 1
        
        chunk = text[start:end]
        chunks.append(chunk)
        
        # 重叠滑动
        start = max(start + 1, end - overlap)  # 防止 start 不增加
        
        # 防止无限循环
        if start >= text_length - 100 or start >= text_length:
            # 添加最后一块
            if start < text_length:
                chunks.append(text[start:text_length])
            break
    
    return chunks

# ==================== 主任务函数 ====================

@celery_app.task(bind=True, base=BaseTaskWithDB, name="tasks.document_tasks.process_large_document")
def process_large_document(self, file_path: str, doc_id: int, user_id: Optional[str] = None):
    """
    处理大文档任务（包含知识图谱提取）
    
    参数:
    - file_path: 文件路径
    - doc_id: 文档ID
    - user_id: 用户ID（可选）
    """
    task_id = self.request.id
    logger.info(f"🚀 开始后台处理文档 ID: {doc_id}, 任务ID: {task_id}")
    
    try:
        # 1. 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 2. 根据文件类型选择处理方式
        file_ext = os.path.splitext(file_path)[1].lower()
        file_size = os.path.getsize(file_path) / (1024 * 1024)  # MB
        
        logger.info(f"📄 处理文件: {file_path} ({file_size:.2f}MB)")
        
        # 3. 读取文件内容
        content = ""
        total_steps = 0
        
        if file_ext == '.pdf':
            content, total_steps = _process_pdf(file_path)
        elif file_ext == '.docx':
            content, total_steps = _process_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            content, total_steps = _process_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        if not content or len(content) < 50:
            raise ValueError(f"文档内容太短或无效: {len(content)} 字符")
        
        # 4. 更新任务进度（20%）
        self.update_state(
            state='PROGRESS',
            meta={
                'current': 1,
                'total': 5,
                'stage': '文件解析完成',
                'progress': 20,
                'details': f"已解析 {len(content)} 字符"
            }
        )
        
        # 5. 分割文档为块
        chunks = _split_into_chunks(content, max_chunk_size=2000)
        logger.info(f"📑 文档分割为 {len(chunks)} 个块")
        
        # 6. 更新任务进度（40%）
        self.update_state(
            state='PROGRESS',
            meta={
                'current': 2,
                'total': 5,
                'stage': '文档分割完成',
                'progress': 40,
                'details': f"已分割为 {len(chunks)} 个块"
            }
        )
        
        # 7. 为每个块生成向量
        vectors = []
        
        for i, chunk in enumerate(chunks):
            try:
                # 根据块的大小动态选择维度
                chunk_len = len(chunk)
                if chunk_len > 1000:
                    mode = "precise"
                elif chunk_len > 500:
                    mode = "balanced"
                else:
                    mode = "fast"
                
                # 生成向量
                vector = asyncio.run(generate_vector(chunk, mode=mode))
                vectors.append(vector)
                
                # 更新子进度
                if i % max(1, len(chunks) // 10) == 0:
                    progress = 40 + (i / len(chunks)) * 20
                    self.update_state(
                        state='PROGRESS',
                        meta={
                            'current': 2 + (i / len(chunks)),
                            'total': 5,
                            'stage': '向量生成中',
                            'progress': progress,
                            'details': f"已生成 {i+1}/{len(chunks)} 个向量"
                        }
                    )
                    
            except Exception as e:
                logger.warning(f"向量生成失败，块 {i+1}: {e}")
                # 使用空向量作为占位符
                vectors.append([0.0] * 1536)  # 假设1536维
        
        # 8. 更新任务进度（60%）
        self.update_state(
            state='PROGRESS',
            meta={
                'current': 3,
                'total': 5,
                'stage': '向量生成完成',
                'progress': 60,
                'details': f"已生成 {len(vectors)} 个向量"
            }
        )
        
        # 9. 存储到数据库
        file_name = os.path.basename(file_path)
        with self.db_session as session:
            # 更新或创建文档记录
            document = session.get(Document, doc_id)
            if not document:
                document = Document(
                    id=doc_id,
                    title=file_name,
                    content=content[:10000] + "..." if len(content) > 10000 else content,
                    embedding=vectors[0] if vectors else None,
                    user_id=user_id,
                    graph_extracted=False
                )
                session.add(document)
            else:
                document.title = file_name
                document.content = content[:10000] + "..." if len(content) > 10000 else content
                document.embedding = vectors[0] if vectors else None
                document.graph_extracted = False
            
            session.commit()
            session.refresh(document)
            
            # 记录向量化结果
            logger.info(f"✅ 文档 {doc_id} 向量化完成，存储 {len(vectors)} 个向量")
            
            # 10. 提取知识图谱
            self.update_state(
                state='PROGRESS',
                meta={
                    'current': 4,
                    'total': 5,
                    'stage': '开始知识图谱提取',
                    'progress': 70,
                    'details': "正在从文档中提取实体和关系..."
                }
            )
            
            try:
                # 创建图谱提取器
                extractor = get_graph_extractor(session)
                
                # 提取图谱（同步调用异步函数）
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    graph_result = loop.run_until_complete(
                        extractor.extract_from_document(document)
                    )
                    
                    if graph_result['success']:
                        logger.info(f"✅ 文档 {doc_id} 知识图谱提取成功: "
                                   f"{graph_result['entities_saved']} 实体, "
                                   f"{graph_result['relations_saved']} 关系")
                        
                        self.update_state(
                            state='PROGRESS',
                            meta={
                                'current': 4.5,
                                'total': 5,
                                'stage': '知识图谱提取完成',
                                'progress': 85,
                                'details': f"提取了 {graph_result['entities_saved']} 个实体和 {graph_result['relations_saved']} 个关系"
                            }
                        )
                    else:
                        logger.warning(f"⚠️ 文档 {doc_id} 图谱提取失败: {graph_result.get('error', '未知错误')}")
                finally:
                    loop.close()
                    
            except Exception as e:
                logger.error(f"❌ 知识图谱提取失败: {e}", exc_info=True)
                # 图谱提取失败不应中断整个任务，只记录日志
        
            # 11. 清理临时文件
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.debug(f"🗑️ 已清理临时文件: {file_path}")
            except Exception as e:
                logger.warning(f"清理临时文件失败: {e}")
            
            # 12. 最终进度更新
            self.update_state(
                state='PROGRESS',
                meta={
                    'current': 5,
                    'total': 5,
                    'stage': '处理完成',
                    'progress': 100,
                    'details': f"文档已成功处理，ID: {doc_id}"
                }
            )
            
            return {
                "status": "completed",
                "doc_id": doc_id,
                "title": file_name,
                "chunks_count": len(chunks),
                "vectors_count": len(vectors),
                "content_length": len(content),
                "graph_extracted": document.graph_extracted,
                "task_id": task_id
            }
    
    except Exception as e:
        logger.error(f"❌ 文档处理失败: {e}", exc_info=True)
        raise self.retry(exc=e, countdown=60)

# ==================== 知识图谱提取任务 ====================

@celery_app.task(bind=True, base=BaseTaskWithDB, name="tasks.document_tasks.extract_graph_from_document")
def extract_graph_from_document(self, document_id: int, user_id: Optional[str] = None):
    """独立的知识图谱提取任务"""
    task_id = self.request.id
    logger.info(f"🧠 开始知识图谱提取，文档ID: {document_id}, 任务ID: {task_id}")
    
    try:
        with self.db_session as session:
            # 获取文档
            document = session.get(Document, document_id)
            
            if not document:
                self.update_state(state='FAILURE', meta={'error': '文档不存在'})
                return {"error": "文档不存在"}
            
            # 如果已经提取过，跳过
            if document.graph_extracted:
                self.update_state(
                    state='SUCCESS',
                    meta={'message': '已提取过', 'doc_id': document_id}
                )
                return {"message": "已提取过", "doc_id": document_id}
            
            # 检查文档内容
            if not document.content or len(document.content) < 50:
                self.update_state(state='FAILURE', meta={'error': '文档内容太短'})
                return {"error": "文档内容太短"}
            
            # 创建提取器
            extractor = get_graph_extractor(session)
            
            # 更新任务进度
            self.update_state(
                state='PROGRESS',
                meta={
                    'current': 1,
                    'total': 4,
                    'stage': '开始提取',
                    'progress': 25,
                    'details': '正在初始化图谱提取器...'
                }
            )
            
            # 提取图谱
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                # 提取实体
                self.update_state(
                    state='PROGRESS',
                    meta={
                        'current': 2,
                        'total': 4,
                        'stage': '提取实体',
                        'progress': 50,
                        'details': '正在从文档中提取实体...'
                    }
                )
                
                graph_result = loop.run_until_complete(
                    extractor.extract_from_document(document)
                )
                
                if graph_result['success']:
                    # 更新文档状态
                    document.graph_extracted = True
                    document.graph_extraction_time = time.time()
                    session.commit()
                    
                    self.update_state(
                        state='SUCCESS',
                        meta={
                            'message': '图谱提取成功',
                            'doc_id': document_id,
                            'entities': graph_result['entities_saved'],
                            'relations': graph_result['relations_saved']
                        }
                    )
                    
                    logger.info(f"✅ 知识图谱提取成功: 文档 {document_id}, "
                               f"实体: {graph_result['entities_saved']}, "
                               f"关系: {graph_result['relations_saved']}")
                    
                    return {
                        "success": True,
                        "doc_id": document_id,
                        "entities": graph_result['entities_saved'],
                        "relations": graph_result['relations_saved'],
                        "entities_list": [e.get('name', '') for e in graph_result.get('entities', [])][:10]
                    }
                else:
                    self.update_state(
                        state='FAILURE',
                        meta={'error': graph_result.get('error', '未知错误')}
                    )
                    return {
                        "success": False,
                        "error": graph_result.get('error', '未知错误')
                    }
                    
            except Exception as e:
                logger.error(f"图谱提取任务失败: {e}", exc_info=True)
                self.update_state(state='FAILURE', meta={'error': str(e)})
                return {"success": False, "error": str(e)}
            finally:
                loop.close()
                
    except Exception as e:
        logger.error(f"知识图谱提取任务失败: {e}", exc_info=True)
        self.update_state(state='FAILURE', meta={'error': str(e)})
        return {"success": False, "error": str(e)}

@celery_app.task(bind=True, base=BaseTaskWithDB, name="tasks.document_tasks.batch_extract_graphs")
def batch_extract_graphs(self, document_ids: List[int], user_id: Optional[str] = None):
    """批量提取知识图谱"""
    task_id = self.request.id
    logger.info(f"🧠 开始批量提取知识图谱，共 {len(document_ids)} 个文档，任务ID: {task_id}")
    
    results = []
    with self.db_session as session:
        for i, doc_id in enumerate(document_ids):
            try:
                # 获取文档
                document = session.get(Document, doc_id)
                if not document:
                    results.append({
                        "doc_id": doc_id,
                        "status": "failed",
                        "error": "文档不存在"
                    })
                    continue
                
                # 如果已经提取过，跳过
                if document.graph_extracted:
                    results.append({
                        "doc_id": doc_id,
                        "status": "skipped",
                        "message": "已提取过"
                    })
                    continue
                
                # 检查文档内容
                if not document.content or len(document.content) < 50:
                    results.append({
                        "doc_id": doc_id,
                        "status": "skipped",
                        "reason": "内容太短"
                    })
                    continue
                
                # 创建提取器
                extractor = get_graph_extractor(session)
                
                # 提取图谱
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    graph_result = loop.run_until_complete(
                        extractor.extract_from_document(document)
                    )
                    
                    if graph_result['success']:
                        document.graph_extracted = True
                        document.graph_extraction_time = time.time()
                        session.commit()
                        
                        results.append({
                            "doc_id": doc_id,
                            "status": "success",
                            "entities": graph_result['entities_saved'],
                            "relations": graph_result['relations_saved']
                        })
                    else:
                        results.append({
                            "doc_id": doc_id,
                            "status": "failed",
                            "error": graph_result.get('error', '未知错误')
                        })
                finally:
                    loop.close()
                
            except Exception as e:
                logger.error(f"文档 {doc_id} 图谱提取失败: {e}")
                results.append({
                    "doc_id": doc_id,
                    "status": "failed",
                    "error": str(e)
                })
            
            # 更新进度
            progress = (i + 1) / len(document_ids) * 100
            self.update_state(
                state='PROGRESS',
                meta={
                    'current': i + 1,
                    'total': len(document_ids),
                    'stage': '提取中',
                    'progress': progress,
                    'details': f"已处理 {i+1}/{len(document_ids)} 个文档"
                }
            )
    
    success_count = len([r for r in results if r['status'] == 'success'])
    failed_count = len([r for r in results if r['status'] == 'failed'])
    skipped_count = len([r for r in results if r['status'] in ['skipped', 'skipped']])
    
    return {
        "status": "completed",
        "total_docs": len(document_ids),
        "success_count": success_count,
        "failed_count": failed_count,
        "skipped_count": skipped_count,
        "results": results,
        "task_id": task_id
    }

@celery_app.task(bind=True, base=BaseTaskWithDB, name="tasks.document_tasks.reprocess_all_graphs")
def reprocess_all_graphs(self, user_id: Optional[str] = None):
    """重新提取所有文档的知识图谱"""
    task_id = self.request.id
    logger.info(f"🔄 开始重新提取所有文档的知识图谱，任务ID: {task_id}")
    
    with self.db_session as session:
        # 获取所有文档
        if user_id:
            documents = session.exec(
                select(Document).where(Document.user_id == user_id)
            ).all()
        else:
            documents = session.exec(select(Document)).all()
        
        logger.info(f"找到 {len(documents)} 个文档")
        
        # 重置所有文档的图谱提取状态
        for doc in documents:
            doc.graph_extracted = False
            doc.graph_extraction_time = None
        
        session.commit()
        
        # 提取图谱
        extractor = get_graph_extractor(session)
        
        results = []
        for i, document in enumerate(documents):
            try:
                # 跳过内容太短的文档
                if not document.content or len(document.content) < 50:
                    results.append({
                        "doc_id": document.id,
                        "title": document.title,
                        "status": "skipped",
                        "reason": "内容太短"
                    })
                    continue
                
                # 提取图谱
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    graph_result = loop.run_until_complete(
                        extractor.extract_from_document(document)
                    )
                    
                    if graph_result['success']:
                        document.graph_extracted = True
                        document.graph_extraction_time = time.time()
                        
                        results.append({
                            "doc_id": document.id,
                            "title": document.title,
                            "status": "success",
                            "entities": graph_result['entities_saved'],
                            "relations": graph_result['relations_saved']
                        })
                    else:
                        results.append({
                            "doc_id": document.id,
                            "title": document.title,
                            "status": "failed",
                            "error": graph_result.get('error', '未知错误')
                        })
                finally:
                    loop.close()
                
            except Exception as e:
                logger.error(f"文档 {document.id} 图谱提取失败: {e}")
                results.append({
                    "doc_id": document.id,
                    "title": document.title,
                    "status": "failed",
                    "error": str(e)
                })
            
            # 更新进度
            progress = (i + 1) / len(documents) * 100
            self.update_state(
                state='PROGRESS',
                meta={
                    'current': i + 1,
                    'total': len(documents),
                    'stage': '重新提取中',
                    'progress': progress,
                    'details': f"已处理 {i+1}/{len(documents)} 个文档"
                }
            )
        
        session.commit()
        
        success_count = len([r for r in results if r['status'] == 'success'])
        failed_count = len([r for r in results if r['status'] == 'failed'])
        skipped_count = len([r for r in results if r['status'] == 'skipped'])
        
        logger.info(f"重新提取完成: 成功 {success_count}, 失败 {failed_count}, 跳过 {skipped_count}")
        
        return {
            "status": "completed",
            "total_docs": len(documents),
            "success_count": success_count,
            "failed_count": failed_count,
            "skipped_count": skipped_count,
            "results": results[:20],  # 只返回前20个结果
            "task_id": task_id
        }

# ==================== 批量处理任务 ====================

@celery_app.task(bind=True, base=BaseTaskWithDB, name="tasks.document_tasks.batch_process_documents")
def batch_process_documents(self, file_paths: list[str], user_id: Optional[str] = None):
    """批量处理多个文档（包含知识图谱提取）"""
    task_id = self.request.id
    logger.info(f"🚀 开始批量处理 {len(file_paths)} 个文档，任务ID: {task_id}")
    
    results = []
    for i, file_path in enumerate(file_paths):
        try:
            # 为每个文档创建一个唯一的文档ID
            doc_id = int(time.time() * 1000) + i  # 使用时间戳避免冲突
            
            # 为每个文档创建一个子任务
            sub_task = process_large_document.apply_async(
                args=[file_path, doc_id, user_id],
                queue="documents"
            )
            
            results.append({
                "file_path": file_path,
                "doc_id": doc_id,
                "task_id": sub_task.id,
                "status": "queued"
            })
            
            # 更新进度
            progress = (i + 1) / len(file_paths) * 100
            self.update_state(
                state='PROGRESS',
                meta={
                    'current': i + 1,
                    'total': len(file_paths),
                    'stage': '任务分发中',
                    'progress': progress,
                    'details': f"已分发 {i+1}/{len(file_paths)} 个任务"
                }
            )
            
        except Exception as e:
            logger.error(f"分发任务失败: {file_path} - {e}")
            results.append({
                "file_path": file_path,
                "error": str(e),
                "status": "failed"
            })
    
    return {
        "status": "completed",
        "total_tasks": len(file_paths),
        "results": results,
        "task_id": task_id
    }

@celery_app.task(bind=True, base=BaseTaskWithDB, name="tasks.document_tasks.cleanup_orphaned_entities")
def cleanup_orphaned_entities(self, user_id: Optional[str] = None):
    """清理孤立的实体（没有关联文档的实体）"""
    task_id = self.request.id
    logger.info(f"🧹 开始清理孤立实体，任务ID: {task_id}")
    
    with self.db_session as session:
        # 获取所有实体
        if user_id:
            entities = session.exec(
                select(Entity).where(Entity.user_id == user_id)
            ).all()
        else:
            entities = session.exec(select(Entity)).all()
        
        logger.info(f"找到 {len(entities)} 个实体")
        
        orphaned_entities = []
        for i, entity in enumerate(entities):
            # 检查实体是否有关联的文档
            if not entity.documents or len(entity.documents) == 0:
                # 检查实体是否有关系
                if (not entity.outgoing_edges or len(entity.outgoing_edges) == 0) and \
                   (not entity.incoming_edges or len(entity.incoming_edges) == 0):
                    orphaned_entities.append(entity)
            
            # 每处理100个实体更新一次进度
            if i % 100 == 0:
                progress = (i + 1) / len(entities) * 100
                self.update_state(
                    state='PROGRESS',
                    meta={
                        'current': i + 1,
                        'total': len(entities),
                        'stage': '扫描中',
                        'progress': progress,
                        'details': f"已扫描 {i+1}/{len(entities)} 个实体，找到 {len(orphaned_entities)} 个孤立实体"
                    }
                )
        
        # 删除孤立实体
        deleted_count = 0
        for entity in orphaned_entities:
            try:
                session.delete(entity)
                deleted_count += 1
            except Exception as e:
                logger.error(f"删除实体失败 {entity.id}: {e}")
        
        session.commit()
        
        logger.info(f"清理完成: 删除 {deleted_count} 个孤立实体")
        
        return {
            "status": "completed",
            "total_entities": len(entities),
            "orphaned_count": len(orphaned_entities),
            "deleted_count": deleted_count,
            "task_id": task_id
        }

@celery_app.task(bind=True, base=BaseTaskWithDB, name="tasks.document_tasks.cleanup_orphaned_edges")
def cleanup_orphaned_edges(self, user_id: Optional[str] = None):
    """清理孤立的关系边（源或目标实体不存在的边）"""
    task_id = self.request.id
    logger.info(f"🧹 开始清理孤立关系边，任务ID: {task_id}")
    
    with self.db_session as session:
        # 获取所有关系边
        if user_id:
            edges = session.exec(
                select(GraphEdge).where(GraphEdge.user_id == user_id)
            ).all()
        else:
            edges = session.exec(select(GraphEdge)).all()
        
        logger.info(f"找到 {len(edges)} 个关系边")
        
        orphaned_edges = []
        for i, edge in enumerate(edges):
            # 检查源实体和目标实体是否存在
            source_exists = session.get(Entity, edge.source_id) is not None
            target_exists = session.get(Entity, edge.target_id) is not None
            
            if not source_exists or not target_exists:
                orphaned_edges.append(edge)
            
            # 每处理100个边更新一次进度
            if i % 100 == 0:
                progress = (i + 1) / len(edges) * 100
                self.update_state(
                    state='PROGRESS',
                    meta={
                        'current': i + 1,
                        'total': len(edges),
                        'stage': '扫描中',
                        'progress': progress,
                        'details': f"已扫描 {i+1}/{len(edges)} 个关系边，找到 {len(orphaned_edges)} 个孤立边"
                    }
                )
        
        # 删除孤立边
        deleted_count = 0
        for edge in orphaned_edges:
            try:
                session.delete(edge)
                deleted_count += 1
            except Exception as e:
                logger.error(f"删除关系边失败 {edge.id}: {e}")
        
        session.commit()
        
        logger.info(f"清理完成: 删除 {deleted_count} 个孤立关系边")
        
        return {
            "status": "completed",
            "total_edges": len(edges),
            "orphaned_count": len(orphaned_edges),
            "deleted_count": deleted_count,
            "task_id": task_id
        }