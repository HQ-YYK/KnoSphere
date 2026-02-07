from fastapi import FastAPI, UploadFile, File, Depends, HTTPException
from fastapi.middleware.cors import CORSMiddleware  # 新增导入
from contextlib import asynccontextmanager
from sqlmodel import Session, text
import os
import io
from services.search import hybrid_search

# 数据库和模型导入
from database import init_db, engine, get_session
from models import User, Document
from services.embedding import generate_vector

from fastapi.responses import StreamingResponse
from services.llm import get_llm_service
from services.search import hybrid_search
import json

@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    # 启动时执行
    print("🚀 启动 KnoSphere API...")

    # 显示当前Embedding配置
    provider = os.getenv("EMBEDDING_PROVIDER", "openai")
    model = os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")
    dim = os.getenv("VECTOR_DIM", "1536")
    print(f"🤖 当前Embedding配置: {provider} / {model} / {dim}维")
    
    with engine.connect() as conn:
        # 激活向量扩展
        conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector;"))
        conn.commit()
    
    init_db()  # 这会创建所有表，包括 User 和 Document
    print("✅ 数据库初始化完成")
    
    yield
    
    # 关闭时执行
    print("👋 关闭 KnoSphere API...")

app = FastAPI(
    title="KnoSphere API",
    description="2026 企业级智能知识库系统",
    version="1.0.0",
    lifespan=lifespan
)

# 添加 CORS 中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 开发环境允许所有源
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有方法
    allow_headers=["*"],  # 允许所有头部
)


@app.get("/")
async def root():
    return {"message": "欢迎使用 KnoSphere API - 2026 企业级智能知识库系统"}

@app.get("/health")
async def health():
    return {
        "status": "healthy", 
        "service": "KnoSphere API",
        "embedding_provider": os.getenv("EMBEDDING_PROVIDER", "openai"),
        "embedding_model": os.getenv("EMBEDDING_MODEL", "text-embedding-3-small"),
        "vector_dimension": os.getenv("VECTOR_DIM", "1536")
    }

@app.get("/documents")
async def list_documents(
    db: Session = Depends(get_session),
    limit: int = 10,
    offset: int = 0
):
    """获取文档列表"""
    documents = db.query(Document).offset(offset).limit(limit).all()
    return {
        "total": db.query(Document).count(),
        "documents": [
            {
                "id": doc.id,
                "title": doc.title,
                "created_at": doc.created_at,
                "content_preview": doc.content[:100] + "..." if len(doc.content) > 100 else doc.content
            }
            for doc in documents
        ]
    }

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...), 
    db: Session = Depends(get_session)
):
    """
    上传文档并生成向量
    支持格式：.txt, .md, .pdf, .docx
    """
    # 1. 验证文件类型
    allowed_extensions = {'.txt', '.md', '.pdf', '.docx'}
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"不支持的文件格式。支持格式: {', '.join(allowed_extensions)}"
        )
    
    # 2. 读取文件内容
    try:
        content = await file.read()
        
        # 处理不同文件类型
        if file_ext == '.pdf':
            # PDF 处理 - 需要额外的依赖
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text()
            except ImportError:
                # 如果未安装 PyPDF2，提示安装
                raise HTTPException(
                    status_code=400,
                    detail="PDF 处理需要安装 PyPDF2。请运行: uv add PyPDF2"
                )
        elif file_ext == '.docx':
            # DOCX 处理
            try:
                import docx
                doc = docx.Document(io.BytesIO(content))
                text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except ImportError:
                raise HTTPException(
                    status_code=400,
                    detail="DOCX 处理需要安装 python-docx。请运行: uv add python-docx"
                )
        else:
            # 文本文件处理
            text_content = content.decode("utf-8")
            
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="文件编码错误，请使用 UTF-8 编码")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件读取失败: {str(e)}")

    # 3. 检查文本长度
    if len(text_content.strip()) == 0:
        raise HTTPException(status_code=400, detail="文件内容为空")
    
    # 4. 生成向量 (AI 核心步骤)
    try:
        vector = await generate_vector(text_content)
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"向量生成失败: {str(e)}。请检查 OPENAI_API_KEY 环境变量"
        )

    # 5. 存储到 PostgreSQL 17
    try:
        new_doc = Document(
            title=file.filename,
            content=text_content,
            embedding=vector  # 存入我们之前定义的 Vector 字段
        )
        
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        
        return {
            "message": "上传成功", 
            "document_id": new_doc.id,
            "title": new_doc.title,
            "vector_dimensions": len(vector) if vector else 0,
            "content_length": len(text_content)
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"数据库存储失败: {str(e)}")


@app.get("/query")
async def query_knowledge_base(
    q: str,
    top_k: int= 15,
    final_k: int = 3,
    db: Session = Depends(get_session)
):
    """
    智能查询知识库
    
    参数：
    - q: 查询问题
    - top_k: 粗排阶段返回的文档数量（默认15）
    - final_k: 精排后最终返回的文档数量（默认3）
    """
    if not q or not q.strip():
        raise HTTPException(status_code=400, detail="请输入问题")
    
    if len(q.strip()) > 1000:
        raise HTTPException(status_code=400, detail="问题过长，请精简到1000字符以内")
    
    try:
        print(f"🔍 开始搜索: {q}")
        results = await hybrid_search(q, db, top_k=top_k, final_k=final_k)
        
        # 格式化返回结果
        formatted_results = []
        for i, doc in enumerate(results):
            formatted_results.append({
                "rank": i + 1,
                "id": doc.get("id"),
                "title": doc.get("title", "无标题"),
                "score": round(doc.get("score", 0) * 100, 2),  # 转换为百分比
                "content_preview": doc.get("content", "")[:200] + "..." if len(doc.get("content", "")) > 200 else doc.get("content", ""),
                "created_at": doc.get("created_at")
            })
        
        return {
            "query": q,
            "total_results": len(results),
            "results": formatted_results
        }
        
    except Exception as e:
        print(f"❌ 搜索失败: {e}")
        raise HTTPException(status_code=500, detail=f"搜索失败: {str(e)}")

@app.get("/documents/{document_id}")
async def get_document(
    document_id: int,
    db: Session = Depends(get_session)
):
    """获取特定文档的详细信息"""
    document = db.get(Document, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    return {
        "id": document.id,
        "title": document.title,
        "content": document.content,
        "created_at": document.created_at,
        "vector_dimensions": len(document.embedding) if document.embedding else 0
    }

@app.get("/search-test")
async def search_test(
    q: str = "什么是人工智能",
    db: Session = Depends(get_session)
):
    """搜索测试端点（用于快速测试）"""
    try:
        results = await hybrid_search(q, db, top_k=5, final_k=3)
        
        # 如果数据库中没有文档，创建一些测试数据
        if not results:
            from services.embedding import generate_vector
            import datetime
            
            # 创建测试文档
            test_docs = [
                {
                    "title": "人工智能简介",
                    "content": "人工智能（AI）是计算机科学的一个分支，旨在创造能够执行通常需要人类智能的任务的机器。"
                },
                {
                    "title": "机器学习基础",
                    "content": "机器学习是人工智能的一个子领域，使计算机能够在没有明确编程的情况下学习和改进。"
                },
                {
                    "title": "深度学习",
                    "content": "深度学习是机器学习的一个分支，使用多层神经网络来模拟人脑的工作方式。"
                }
            ]
            
            for doc_data in test_docs:
                vector = await generate_vector(doc_data["content"])
                new_doc = Document(
                    title=doc_data["title"],
                    content=doc_data["content"],
                    embedding=vector,
                    created_at=datetime.datetime.utcnow()
                )
                db.add(new_doc)
            
            db.commit()
            
            # 重新搜索
            results = await hybrid_search(q, db, top_k=5, final_k=3)
        
        return {
            "query": q,
            "results": results,
            "message": "测试成功" if results else "数据库为空，已创建测试数据"
        }
        
    except Exception as e:
        return {"error": str(e)}


@app.post("/chat")
async def chat(
    request: dict,
    db: Session = Depends(get_session)
):
    """
    聊天接口 - 流式响应
    
    请求体：
    {
        "query": "用户的问题",
        "top_k": 10,  # 可选，检索文档数量
        "final_k": 3   # 可选，最终使用文档数量
    }
    """
    query = request.get("query", "").strip()
    top_k = request.get("top_k", 10)
    final_k = request.get("final_k", 3)
    
    if not query:
        return StreamingResponse(
            iter(["❌ 请输入问题"]),
            media_type="text/plain"
        )
    
    # 1. 检索最相关的知识
    try:
        docs = await hybrid_search(query, db, top_k=top_k, final_k=final_k)
        
        if not docs:
            return StreamingResponse(
                iter(["🔍 知识库中没有找到相关文档。请先上传一些文档，或者换个问题试试。"]),
                media_type="text/plain"
            )
        
        # 2. 拼接上下文
        context_parts = []
        for i, doc in enumerate(docs[:final_k]):  # 使用 final_k 限制
            context_parts.append(f"【文档{i+1}】{doc.get('title', '无标题')}")
            content_preview = doc.get('content', '')[:500] + "..." if len(doc.get('content', '')) > 500 else doc.get('content', '')
            context_parts.append(f"内容：{content_preview}")
            context_parts.append(f"相关度：{doc.get('score', 0):.2%}")
            context_parts.append("---")
        
        context_text = "\n".join(context_parts)
        
        # 3. 获取 LLM 服务
        llm_service = get_llm_service()
        
        # 4. 返回流式响应
        async def generate():
            # 先返回检索结果摘要
            yield f"🔍 已为您检索到 {len(docs)} 篇相关文档，使用前 {min(final_k, len(docs))} 篇生成回答：\n\n"
            await asyncio.sleep(0.1)
            
            # 然后流式返回 AI 回答
            async for chunk in llm_service.stream_response(query, context_text):
                yield chunk
        
        return StreamingResponse(
            generate(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no"  # 禁用 Nginx 缓冲
            }
        )
        
    except Exception as e:
        print(f"❌ 聊天处理失败: {e}")
        return StreamingResponse(
            iter([f"❌ 处理请求时出错：{str(e)}"]),
            media_type="text/plain"
        )

@app.get("/chat-test")
async def chat_test(
    query: str = "什么是人工智能？",
    db: Session = Depends(get_session)
):
    """聊天测试端点（非流式，用于快速测试）"""
    if not query:
        return {"error": "请输入问题"}
    
    try:
        # 检索文档
        docs = await hybrid_search(query, db, top_k=5, final_k=2)
        
        if not docs:
            return {"answer": "知识库中没有找到相关文档。", "documents": []}
        
        # 构建上下文
        context_parts = []
        for i, doc in enumerate(docs):
            context_parts.append(f"【文档{i+1}】{doc.get('title', '无标题')}")
            context_parts.append(f"内容：{doc.get('content', '')[:300]}...")
            context_parts.append("---")
        
        context_text = "\n".join(context_parts)
        
        # 获取 LLM 服务
        llm_service = get_llm_service()
        
        # 收集流式响应
        full_response = ""
        async for chunk in llm_service.stream_response(query, context_text):
            full_response += chunk
        
        return {
            "query": query,
            "answer": full_response,
            "documents_used": [
                {
                    "title": doc.get("title"),
                    "score": f"{doc.get('score', 0):.2%}",
                    "content_preview": doc.get("content", "")[:100] + "..."
                }
                for doc in docs[:2]
            ]
        }
        
    except Exception as e:
        return {"error": str(e)}

# 如果需要添加更多文件格式处理，可以取消下面的注释并安装相应依赖
# @app.on_event("startup")
# async def check_dependencies():
#     """检查可选依赖"""
#     try:
#         import PyPDF2
#         print("✅ PyPDF2 已安装，支持 PDF 处理")
#     except ImportError:
#         print("⚠️  PyPDF2 未安装，PDF 文件处理将不可用")
#     
#     try:
#         import docx
#         print("✅ python-docx 已安装，支持 DOCX 处理")
#     except ImportError:
#         print("⚠️  python-docx 未安装，DOCX 文件处理将不可用")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8000,
    )